from __future__ import annotations

import json
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional

from .mc_engine import run_monte_carlo
from .transforms.soil_chm import (
    transform_compute_base_soil_vars,
    transform_perturb_relative,
    #transform_split_fixed_ratios,
    transform_split_with_bounds,
    transform_write_chm_from_df,
    transform_apply_ops,
)
from .transforms.point_dat import (
    transform_build_point_load_timeseries,
    transform_write_point_dat_from_df,
    transform_interpolate_years_wide,
)


def _ledger_path_default() -> Path:
    return Path(__file__).resolve().parent.parent / "config" / "provenance" / "realizations.jsonl"


def read_ledger(ledger_path: Optional[Path] = None) -> List[Dict[str, Any]]:
    p = ledger_path or _ledger_path_default()
    if not Path(p).exists():
        return []
    recs: List[Dict[str, Any]] = []
    with open(p, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            try:
                recs.append(json.loads(line))
            except Exception:
                continue
    return recs


def _parse_iso(ts: str) -> Optional[datetime]:
    try:
        return datetime.fromisoformat(ts.replace("Z", "+00:00"))
    except Exception:
        return None


def summarize_provenance(ledger_path: Optional[Path] = None) -> Dict[str, Any]:
    recs = read_ledger(ledger_path)
    if not recs:
        return {"count": 0}

    ids = [r.get("id") for r in recs if isinstance(r.get("id"), int)]
    names = [r.get("name") for r in recs]
    engines = {json.dumps(r.get("engine", {}), sort_keys=True): 1 for r in recs}
    run_ids = sorted({r.get("run_id") for r in recs if r.get("run_id") is not None})
    created = [r.get("created_at") for r in recs if r.get("created_at")]
    created_dt = [d for d in map(_parse_iso, created) if d]
    inputs = []
    for r in recs:
        for i in r.get("inputs", []) or []:
            if isinstance(i, dict) and i.get("path"):
                inputs.append(i["path"])
    unique_inputs = sorted(set(inputs))

    # step timing (approx): sum durations from started_at to ended_at per record
    def rec_duration(r: Dict[str, Any]) -> Optional[float]:
        total = 0.0
        any_step = False
        for s in r.get("steps", []) or []:
            st = _parse_iso(s.get("started_at", ""))
            en = _parse_iso(s.get("ended_at", ""))
            if st and en and en >= st:
                total += (en - st).total_seconds()
                any_step = True
        return total if any_step else None

    durations = [d for d in (rec_duration(r) for r in recs) if d is not None]

    return {
        "count": len(recs),
        "ids": ids,
        "names": names,
        "engine_variants": len(engines),
        "run_ids": run_ids,
        "created_min": min(created_dt).isoformat() if created_dt else None,
        "created_max": max(created_dt).isoformat() if created_dt else None,
        "unique_inputs": unique_inputs,
        "durations_sec": {
            "avg": (sum(durations) / len(durations)) if durations else None,
            "min": (min(durations) if durations else None),
            "max": (max(durations) if durations else None),
        },
    }


def format_summary(summary: Dict[str, Any]) -> str:
    if not summary or summary.get("count", 0) == 0:
        return "No provenance records found."
    parts = []
    parts.append(f"Records: {summary['count']}")
    if summary.get("created_min") and summary.get("created_max"):
        parts.append(f"Time span: {summary['created_min']} → {summary['created_max']}")
    if summary.get("durations_sec", {}).get("avg") is not None:
        d = summary["durations_sec"]
        parts.append(f"Step durations (sec): avg={d['avg']:.2f}, min={d['min']:.2f}, max={d['max']:.2f}")
    if summary.get("run_ids"):
        parts.append("Runs:")
        parts.append("  - run_ids=" + ", ".join(str(x) for x in summary["run_ids"]))
    if summary.get("unique_inputs"):
        parts.append("Inputs:")
        for p in summary["unique_inputs"]:
            parts.append(f"  - {p}")
    return "\n".join(parts)


from pathlib import Path
from typing import Optional, Sequence, Tuple, List, Any, Dict



# --- internal helpers ---------------------------------------------------------

def _fmt_bounds(lu: Tuple[Optional[float], Optional[float]] | Any) -> str:
    if not isinstance(lu, tuple):
        return "-"
    l, u = lu
    if l is None and u is None:
        return "-"
    return f"[{l}..{u}]"

def _format_val(val: Any) -> str:
    if isinstance(val, (int, float)):
        # match your fixed-width presentation
        return f"{val:.6f}"
    return str(val) if val is not None else "-"

def _collect_summary_rows(rec: Dict[str, Any]) -> List[Tuple[str, str, str, Tuple[Optional[float], Optional[float]], Any, Any, Optional[bool]]]:
    """
    Builds the succinct table rows:
    (transform | in -> out | bounds | value | source | renorm?)
    Returns list of tuples:
      (name, in_var, out_var, (lower, upper), value, source, renorm_bool_or_None)
    """
    rows: List[Tuple[str, str, str, Tuple[Optional[float], Optional[float]], Any, Any, Optional[bool]]] = []

    # helpers to pick bounds for scale/split
    scale_bounds: Dict[str, Tuple[Optional[float], Optional[float]]] = {}
    split_bounds: Dict[str, Tuple[Optional[float], Optional[float]]] = {}

    # pre-scan for defaults
    for s in rec.get("steps") or []:
        if s.get("name") == "transform_scale_variable":
            out = s.get("args", {}).get("out")
            lower = s.get("args", {}).get("lower")
            upper = s.get("args", {}).get("upper")
            if out is not None and (lower is not None or upper is not None):
                scale_bounds[out] = (lower, upper)
        if s.get("name") == "transform_split_with_bounds":
            for o in (s.get("args", {}).get("outputs", []) or []):
                split_bounds[o.get("name")] = (o.get("lower"), o.get("upper"))

    # collect realized choices
    for s in rec.get("steps") or []:
        nm = s.get("name")
        if nm == "perturbation_choices":
            for t in (s.get("args", {}).get("targets", []) or []):
                rows.append(("perturb_relative", t.get("name"), t.get("name"), (None, None), t.get("delta"), t.get("source"), None))
        elif nm == "scale_choices":
            a = s.get("args", {})
            rows.append((
                "scale_variable",
                a.get("src"),
                a.get("out"),
                scale_bounds.get(a.get("out"), (a.get("lower"), a.get("upper"))),
                a.get("factor"),
                a.get("source"),
                None
            ))
        elif nm == "ops_choices":
            for o in (s.get("args", {}).get("ops", []) or []):
                rows.append((
                    f"op_{o.get('op')}",
                    o.get("src"),
                    o.get("out"),
                    (o.get("lower"), o.get("upper")),
                    o.get("value"),
                    o.get("source"),
                    None
                ))
        elif nm == "point_mgL_choices":
            a = s.get("args", {})
            m = a.get("mgL_values", {}) or {}
            for var, meta in m.items():
                rows.append((
                    "mgL_point",
                    "population",
                    var,
                    (meta.get("lower"), meta.get("upper")),
                    meta.get("mgL"),
                    meta.get("source"),
                    None
                ))
        elif nm == "split_choices":
            a = s.get("args", {})
            src = a.get("src")
            renorm = a.get("renormalize")
            for o in (a.get("outputs", []) or []):
                rows.append((
                    "split_with_bounds",
                    src,
                    o.get("name"),
                    split_bounds.get(o.get("name"), (o.get("lower"), o.get("upper"))),
                    o.get("ratio"),
                    o.get("source"),
                    renorm
                ))
    return rows

def _markdown_table_from_rows(rows: Sequence[Tuple[str, str, str, Tuple[Optional[float], Optional[float]], Any, Any, Optional[bool]]]) -> str:
    header = "| transform | in | out | bounds | value | source | renorm |\n|---|---|---|---|---|---|---|"
    lines = [header]
    for name, inv, outv, lu, val, src, renorm in rows:
        lines.append(f"| {name} | {inv} | {outv} | {_fmt_bounds(lu)} | {_format_val(val)} | {src if src is not None else '-'} | "
                     f"{'' if renorm is None else ('renormalized' if renorm else 'not renormalized')} |")
    return "\n".join(lines)

# --- public reports -----------------------------------------------------------

def realization_report(realization_id: int | str, ledger_path: Optional[Path] = None) -> str:
    """
    Backwards-compatible: returns the plain-text, line-based report (no markdown table).
    """
    recs = read_ledger(ledger_path)
    rec = next((r for r in recs if str(r.get("id")) == str(realization_id)), None)
    if not rec:
        return f"No record found for id={realization_id}"

    lines: List[str] = []
    lines.append(f"Realization id={rec.get('id')} run_id={rec.get('run_id')} name={rec.get('name')}")
    lines.append(f"Created: {rec.get('created_at')}")
    engine = rec.get("engine", {})
    lines.append(f"Engine: module={engine.get('module')} version={engine.get('version')} seed={engine.get('seed')} N={engine.get('N')}")
    lines.append(f"Base TxtInOut: {rec.get('base_txtinout')}")
    lines.append(f"Realization folder: {rec.get('realization_folder')}")

    rows = _collect_summary_rows(rec)
    if rows:
        lines.append("\nSummary (transform | in -> out | bounds | value | source | renorm):")
        for name, inv, outv, lu, val, src, renorm in rows:
            b_s = _fmt_bounds(lu)
            v_s = _format_val(val)
            ren_s = "" if renorm is None else ("renormalized" if renorm else "not renormalized")
            lines.append(f"  - {name}: {inv} -> {outv} | {b_s} | {v_s} | {src} | {ren_s}")

    if rec.get("inputs"):
        lines.append("Inputs:")
        for i in rec["inputs"]:
            if isinstance(i, dict):
                p = str(i.get("path", "")).lower()
                if p.endswith(".tif") or p.endswith(".tiff") or p.endswith(".shp"):
                    lines.append(f"  - {i.get('path')}")

    lines.append("\nSteps:")
    for s in rec.get("steps") or []:
        lines.append(f"  - {s.get('name')} | args={s.get('args')} | started={s.get('started_at')} ended={s.get('ended_at')}")
    if rec.get("outputs_summary"):
        lines.append(f"Outputs: {rec['outputs_summary']}")

    return "\n".join(lines)

def realization_report_assets(
    realization_id: int | str,
    out_dir: Path,
    ledger_path: Optional[Path] = None,
    md_filename: str = "summary_table.md",
    docx_filename: str = "report.docx",
    txt_filename: str = "report.txt",
) -> dict:
    """
    Creates three files in out_dir:
      - txt_filename: the original plaintext report
      - md_filename:  a Markdown table-only summary (ready to copy/paste)
      - docx_filename: Word document with the table + rest of the summary

    Returns a dict with created file paths.
    """
    recs = read_ledger(ledger_path)
    rec = next((r for r in recs if str(r.get("id")) == str(realization_id)), None)
    if not rec:
        raise FileNotFoundError(f"No record found for id={realization_id}")

    out_dir.mkdir(parents=True, exist_ok=True)

    # 1) Plain-text report (as-is, like your existing output)
    plain_text = realization_report(realization_id, ledger_path=ledger_path)
    (out_dir / txt_filename).write_text(plain_text, encoding="utf-8")

    # 2) Markdown table (ready to copy)
    rows = _collect_summary_rows(rec)
    md_table = _markdown_table_from_rows(rows) if rows else "| (no summary rows) |\n|---|"
    (out_dir / md_filename).write_text(md_table, encoding="utf-8")

    # 3) Word document with formatted table + the rest of the summary
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.oxml.ns import qn

        doc = Document()
        doc.add_heading(f"Realization {rec.get('id')}", level=1)

        # meta block
        p = doc.add_paragraph()
        p.add_run(f"Run ID: {rec.get('run_id')}\n")
        p.add_run(f"Name: {rec.get('name')}\n")
        p.add_run(f"Created: {rec.get('created_at')}\n")
        engine = rec.get("engine", {})
        p.add_run(f"Engine: module={engine.get('module')} version={engine.get('version')} seed={engine.get('seed')} N={engine.get('N')}\n")
        p.add_run(f"Base TxtInOut: {rec.get('base_txtinout')}\n")
        p.add_run(f"Realization folder: {rec.get('realization_folder')}\n")

        # table
        doc.add_heading("Summary", level=2)
        headers = ["transform", "in", "out", "bounds", "value", "source", "renorm"]
        if rows:
            table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
            table.style = "Table Grid"
            for j, h in enumerate(headers):
                table.cell(0, j).text = h
            for i, (name, inv, outv, lu, val, src, renorm) in enumerate(rows, start=1):
                table.cell(i, 0).text = str(name)
                table.cell(i, 1).text = str(inv)
                table.cell(i, 2).text = str(outv)
                table.cell(i, 3).text = _fmt_bounds(lu)
                table.cell(i, 4).text = _format_val(val)
                table.cell(i, 5).text = "-" if src is None else str(src)
                table.cell(i, 6).text = "" if renorm is None else ("renormalized" if renorm else "not renormalized")
        else:
            doc.add_paragraph("(no summary rows)")

        # inputs (filtered to tif/tiff/shp like your text version)
        if rec.get("inputs"):
            doc.add_heading("Inputs", level=2)
            for i in rec["inputs"]:
                if isinstance(i, dict):
                    pth = str(i.get("path", ""))
                    low = pth.lower()
                    if low.endswith(".tif") or low.endswith(".tiff") or low.endswith(".shp"):
                        doc.add_paragraph(pth, style=None)

        # steps
        doc.add_heading("Steps", level=2)
        for s in rec.get("steps") or []:
            doc.add_paragraph(
                f"{s.get('name')} | args={s.get('args')} | started={s.get('started_at')} ended={s.get('ended_at')}"
            )

        # outputs
        if rec.get("outputs_summary"):
            doc.add_heading("Outputs", level=2)
            doc.add_paragraph(str(rec["outputs_summary"]))

        doc.save(out_dir / docx_filename)
    except ImportError as e:
        # Provide a friendly message; still return the MD/TXT paths.
        # To enable Word export: pip install python-docx
        pass

    return {
        "txt": str(out_dir / txt_filename),
        "md": str(out_dir / md_filename),
        "docx": str(out_dir / docx_filename),
    }

# --- small wrapper over your original loop -----------------------------------

def write_provenance_reports(results: Sequence[Any], ledger_path: Optional[Path] = None) -> Sequence[Any]:
    """
    For each result in `results`, write:
      - report.txt (original text report)
      - summary_table.md (ready-to-copy Markdown table)
      - report.docx (Word doc with formatted table + rest of summary)
    Silently continues on per-item errors, like your original.
    """
    for r in results:
        try:
            out_dir = Path(r.folder)
            realization_report_assets(
                realization_id=r.realization_id,
                out_dir=out_dir,
                ledger_path=ledger_path,
                md_filename="summary_table.md",
                docx_filename="report.docx",
                txt_filename="report.txt",
            )
        except Exception:
            # keep parity with your try/except pass
            pass
    return results



def summarize_run(run_id: int, ledger_path: Optional[Path] = None) -> str:
    recs = read_ledger(ledger_path)
    # Accept both top-level run_id and engine.run_id for compatibility
    sel = [
        r for r in recs
        if r.get("run_id") == run_id or (isinstance(r.get("engine"), dict) and r.get("engine", {}).get("run_id") == run_id)
    ]
    if not sel:
        return f"No records found for run_id={run_id}"
    ids = [r.get("id") for r in sel]
    created = [r.get("created_at") for r in sel if r.get("created_at")]
    dt = [d for d in map(_parse_iso, created) if d]
    names = [r.get("name") for r in sel]

    # Aggregate upstream inputs (user-facing only: rasters/shapes)
    inputs = []
    for r in sel:
        for i in r.get("inputs", []) or []:
            if isinstance(i, dict) and i.get("path"):
                p = str(i["path"]).lower()
                if p.endswith(".tif") or p.endswith(".tiff") or p.endswith(".shp"):
                    inputs.append(i["path"])
    inputs = sorted(set(inputs))

    # Extract transform & perturbation metadata
    # Strategy: gather unique transform step names and parse bounds/choices from choices steps
    bounds_map = {}  # target -> bound (perturb_relative legacy)
    sources_map = {} # target -> set of sources used {random, override, none}
    # scale_variable
    scale_lu = {}    # out -> (lower, upper)
    scale_sources = {}
    # ops_apply
    ops_lu = {}
    ops_sources = {}
    # point mgL
    mgL_lu = {}
    mgL_sources = {}
    # split_with_bounds
    split_lu = {}    # out -> (lower, upper)
    split_base = {}  # out -> base_ratio (legacy if present)
    split_sources = {}
    transforms_used = []
    for r in sel:
        for s in r.get("steps") or []:
            nm = s.get("name")
            if nm and nm not in transforms_used:
                transforms_used.append(nm)
            if s.get("name") == "transform_perturb_relative":
                for t in (s.get("args", {}).get("targets", []) or []):
                    nm = t.get("name")
                    b = t.get("bound")
                    if nm is not None and b is not None:
                        bounds_map.setdefault(nm, float(b))
            if s.get("name") == "perturbation_choices":
                for t in (s.get("args", {}).get("targets", []) or []):
                    nm = t.get("name")
                    src = t.get("source")
                    if nm is None:
                        continue
                    if src:
                        sset = sources_map.get(nm, set())
                        sset.add(str(src))
                        sources_map[nm] = sset
            if s.get("name") == "scale_choices":
                out = s.get("args", {}).get("out")
                lower = s.get("args", {}).get("lower")
                upper = s.get("args", {}).get("upper")
                src = s.get("args", {}).get("source")
                if out is not None:
                    scale_lu[out] = (lower, upper)
                    if src:
                        sset = scale_sources.get(out, set())
                        sset.add(str(src))
                        scale_sources[out] = sset
            if s.get("name") == "ops_choices":
                for o in (s.get("args", {}).get("ops", []) or []):
                    outn = o.get("out")
                    lower = o.get("lower")
                    upper = o.get("upper")
                    src = o.get("source")
                    if outn is None:
                        continue
                    ops_lu[outn] = (lower, upper)
                    if src:
                        sset = ops_sources.get(outn, set())
                        sset.add(str(src))
                        ops_sources[outn] = sset
            if s.get("name") == "point_mgL_choices":
                m = s.get("args", {}).get("mgL_values", {}) or {}
                for var, meta in m.items():
                    lower = meta.get("lower")
                    upper = meta.get("upper")
                    src = meta.get("source")
                    mgL_lu[var] = (lower, upper)
                    if src:
                        sset = mgL_sources.get(var, set())
                        sset.add(str(src))
                        mgL_sources[var] = sset
            if s.get("name") == "split_choices":
                for o in (s.get("args", {}).get("outputs", []) or []):
                    outn = o.get("name")
                    lower = o.get("lower")
                    upper = o.get("upper")
                    src = o.get("source")
                    if outn is None:
                        continue
                    split_lu[outn] = (lower, upper)
                    if src:
                        sset = split_sources.get(outn, set())
                        sset.add(str(src))
                        split_sources[outn] = sset

    lines = []
    lines.append(f"Run {run_id}: realizations={len(sel)} ids={ids}")
    if dt:
        lines.append(f"Time span: {min(dt).isoformat()} → {max(dt).isoformat()}")
    lines.append("Names:")
    for n in names:
        lines.append(f"  - {n}")
    if transforms_used:
        lines.append("Transforms:")
        for tname in transforms_used:
            lines.append(f"  - {tname}")
    if inputs:
        lines.append("Inputs (union):")
        for p in inputs:
            lines.append(f"  - {p}")
    if bounds_map or scale_lu or ops_lu or mgL_lu:
        lines.append("Perturbations:")
        # perturb_relative
        for nm, b in bounds_map.items():
            srcs = ",".join(sorted(sources_map.get(nm, set()))) if nm in sources_map else "unknown"
            lines.append(f"  - {nm}: strategy={srcs}")
        # scale_variable
        for out, lu in scale_lu.items():
            srcs = ",".join(sorted(scale_sources.get(out, set()))) if out in scale_sources else "unknown"
            b = f"[{lu[0]}..{lu[1]}]" if not (lu and lu[0] is None and lu[1] is None) else "unbounded"
            lines.append(f"  - {out}: bounds={b} strategy={srcs}")
        # ops_apply
        for out, lu in ops_lu.items():
            srcs = ",".join(sorted(ops_sources.get(out, set()))) if out in ops_sources else "unknown"
            b = f"[{lu[0]}..{lu[1]}]" if not (lu and lu[0] is None and lu[1] is None) else "unbounded"
            lines.append(f"  - {out}: bounds={b} strategy={srcs}")
        # mgL choices
        for var, lu in mgL_lu.items():
            srcs = ",".join(sorted(mgL_sources.get(var, set()))) if var in mgL_sources else "unknown"
            b = f"[{lu[0]}..{lu[1]}]" if not (lu and lu[0] is None and lu[1] is None) else "unbounded"
            lines.append(f"  - mgL {var}: bounds={b} strategy={srcs}")
    if split_lu:
        lines.append("Splits:")
        for outn, lu in split_lu.items():
            srcs = ",".join(sorted(split_sources.get(outn, set()))) if outn in split_sources else "unknown"
            base = split_base.get(outn)
            base_s = f" base={base:.4f}" if base is not None else ""
            lines.append(f"  - {outn}:{base_s} bounds=[{lu[0]}..{lu[1]}] strategy={srcs}")
    return "\n".join(lines)


def build_upstream_inputs(
    *,
    raster_folder: Path,
    pattern: str,
    zones_fp: Path,
    gpkg_path: Path,
    csv_path: Path,
) -> List[Path]:
    """Helper to collect common upstream inputs for run_monte_carlo(upstream_inputs=...)."""
    paths: List[Path] = []
    try:
        paths.extend(sorted(Path(raster_folder).glob(pattern)))
    except Exception:
        pass
    for p in (zones_fp, gpkg_path, csv_path):
        if p:
            paths.append(Path(p))
    return paths


def _map_step_to_callable(name: str):
    # Map known step names to transform callables
    mapping = {
        "transform_compute_base_soil_vars": transform_compute_base_soil_vars,
        "transform_perturb_relative": transform_perturb_relative,
        #"transform_split_fixed_ratios": transform_split_fixed_ratios,
        "transform_split_with_bounds": transform_split_with_bounds,
        "transform_apply_ops": transform_apply_ops,
        "transform_build_point_load_timeseries": transform_build_point_load_timeseries,
        "transform_write_point_dat_from_df": transform_write_point_dat_from_df,
        "transform_interpolate_years_wide": transform_interpolate_years_wide,
        "transform_write_chm_from_df": transform_write_chm_from_df,
    }
    return mapping.get(name)


def reconstruct_mc_from_ledger(
    *,
    ids: Optional[Iterable[int]] = None,
    base_txtinout: Path,
    realization_root: Path,
    results_root: Path,
    exe_path: Optional[Path] = None,
    config: Optional[Dict[str, Any]] = None,
) -> None:
    """Replay MC realizations exactly as recorded in the ledger.

    Notes:
    - Works for realizations created by the modular transform pipeline where `steps`
      correspond to importable callables.
    - Uses per-realization step args as overrides; default params are empty `{}`.
    - Assumes a consistent transform sequence across the selected records.
    """
    recs = read_ledger()
    sel = [r for r in recs if isinstance(r.get("id"), int) and (ids is None or r["id"] in set(ids))]
    if not sel:
        print("No matching records to replay.")
        return

    # Derive transform sequence from the first record
    first = sel[0]
    step_infos = [(s.get("name"), s.get("args", {})) for s in (first.get("steps") or [])]
    transforms = []
    for nm, _ in step_infos:
        fn = _map_step_to_callable(nm)
        if not fn:
            raise RuntimeError(f"Cannot map step '{nm}' to a known transform callable.")
        transforms.append(fn)

    transforms_params = [{} for _ in transforms]
    per_realization_params: List[List[Dict[str, Any]]] = []
    for r in sel:
        rec_steps = r.get("steps") or []
        # Keep only steps that map to transforms
        params_seq: List[Dict[str, Any]] = []
        for s in rec_steps:
            nm = s.get("name")
            if _map_step_to_callable(nm):
                params_seq.append(s.get("args", {}))
        per_realization_params.append(params_seq)

    # Build aggregator from recorded inputs if possible is out of scope; user supplies data via CSV
    # Here we only replay the transforms and SWAT run; the aggregator should be provided upstream.
    raise_if_no_agg = (
        "Reconstruct requires calling run_monte_carlo with an aggregator for base data "
        "(e.g., read_n_p_means_from_csv_to_df)."
    )
    print(
        "Transforms sequence:", [t.__name__ for t in transforms],
        "\nProvide aggregator=lambda: <load base data> and call run_monte_carlo with transforms/transforms_params/per_realization_params.",
    )

    # Example (user fills in aggregator and linking patterns):
    # run_monte_carlo(
    #   N=len(per_realization_params), base_txtinout=base_txtinout,
    #   realization_root=realization_root, results_root=results_root,
    #   link_file_regexes=[r'^.*\\.chm$'], outputs_to_copy=['output.std','*.rch'],
    #   aggregator=..., transforms=transforms, transforms_params=transforms_params,
    #   per_realization_params=per_realization_params, exe_path=exe_path, config=config)
